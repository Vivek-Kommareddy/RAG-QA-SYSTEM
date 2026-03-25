"""Document loader for various file formats.

This module defines a `Document` dataclass to encapsulate loaded text and
associated metadata.  The `load_file` function dispatches based on file
extension and returns a list of documents ready for chunking.
"""

from __future__ import annotations

import datetime as _dt
import os
from dataclasses import dataclass
from typing import List, Dict

import fitz  # PyMuPDF
import docx  # python-docx


@dataclass
class Document:
    """A loaded document with content and metadata.

    The `content` field contains the raw text of the document or chunk.  The
    `metadata` dict holds arbitrary key/value pairs such as filename,
    page_number, file_type, ingested_at timestamp and any other relevant
    attributes.  Downstream components should treat `metadata` as an opaque
    dictionary.
    """

    content: str
    metadata: Dict[str, object]


def _load_pdf(file_path: str) -> List[Document]:
    """Load a PDF document and return a list of per‑page Documents.

    Each page in the PDF is turned into a separate `Document` with its own
    metadata, preserving the page number.  The entire file name and type are
    recorded for later reference.
    """
    docs: List[Document] = []
    with fitz.open(file_path) as pdf:
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text("text")
            metadata = {
                "filename": os.path.basename(file_path),
                "file_type": "pdf",
                "page_number": page_num + 1,
                "ingested_at": _dt.datetime.utcnow().isoformat(),
            }
            docs.append(Document(content=text.strip(), metadata=metadata))
    return docs


def _load_docx(file_path: str) -> List[Document]:
    """Load a DOCX file into a single `Document`.

    The paragraphs are concatenated with newlines.  Only one Document is
    returned because Word files do not have pages in the same sense as PDFs.
    """
    doc = docx.Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs]
    text = "\n".join(paragraphs)
    metadata = {
        "filename": os.path.basename(file_path),
        "file_type": "docx",
        "ingested_at": _dt.datetime.utcnow().isoformat(),
    }
    return [Document(content=text.strip(), metadata=metadata)]


def _load_text(file_path: str, file_type: str) -> List[Document]:
    """Load a plain text or Markdown file.

    The entire file contents become the content of a single Document.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    metadata = {
        "filename": os.path.basename(file_path),
        "file_type": file_type,
        "ingested_at": _dt.datetime.utcnow().isoformat(),
    }
    return [Document(content=text.strip(), metadata=metadata)]


def load_file(file_path: str) -> List[Document]:
    """Load a file and return a list of `Document` objects.

    Dispatch is based on the file extension.  Unsupported file types raise
    a `ValueError`.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _load_pdf(file_path)
    if ext in {".txt", ".md"}:
        return _load_text(file_path, ext.lstrip("."))
    if ext == ".docx":
        return _load_docx(file_path)
    raise ValueError(f"Unsupported file type: {ext}")
